import os
import io
import re
import json
import requests
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

CONFIG_FILE = "config_variaveis.json"

DEFAULT_VARS = [
    {"nome": "Aptidão", "tipo": "codigo", "opcoes": ["I. Lavoura – Aptidão Boa", "IV. Pastagem Plantada", "V. Silvicultura ou Pastagem Natural"]},
    {"nome": "Acesso", "tipo": "codigo", "opcoes": ["Favorável", "Desfavorável", "Regular", "Má"]},
    {"nome": "Nota Agronômica", "tipo": "numero", "opcoes": []},
    {"nome": "Benfeitoria", "tipo": "numero", "opcoes": []}
]

def converter_para_float(texto):
    """Converte strings financeiras e de área (com pontos e vírgulas) de forma segura."""
    if texto is None:
        return 0.0
    s = str(texto).replace("R$", "").replace("m²", "").replace("ha", "").strip()
    if not s:
        return 0.0
    
    # Se contém tanto ponto quanto vírgula (ex: 27.500.000,00)
    if "." in s and "," in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        # Padrão brasileiro decimal (ex: 476,7400)
        s = s.replace(",", ".")
    elif "." in s:
        partes = s.split(".")
        # Múltiplos pontos são separadores de milhar (ex: 27.500.000)
        if len(partes) > 2:
            s = s.replace(".", "")
        elif len(partes) == 2 and len(partes[1]) == 3 and len(partes[0]) <= 3:
            # Caso ambíguo de milhar sem centavos (ex: 27.500)
            s = s.replace(".", "")
    return float(s)

def formatar_moeda_br(valor):
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def formatar_numero_br(valor, casas=2):
    fmt = f"{{:,.{casas}f}}"
    return fmt.format(valor).replace(",", "X").replace(".", ",").replace("X", ".")

def limpar_sufixo_coord(coord_str):
    s = str(coord_str).strip()
    s = re.sub(r'\s*m\s*[ESes]\s*$', '', s, flags=re.IGNORECASE)
    s = re.sub(r'\s*[ESes]\s*$', '', s, flags=re.IGNORECASE)
    s = re.sub(r'\s*m\s*$', '', s, flags=re.IGNORECASE)
    return s.strip()

class AppPesquisaMercado:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Pesquisa de Mercado - SISDEA / Word")
        self.root.geometry("1060x800")

        self.dados_pesquisas = []
        self.item_em_edicao = None
        self.variaveis_config = self._carregar_config_variaveis()

        self._criar_menu()
        self._criar_layout()
        self._atualizar_interface_variaveis()

    def _carregar_config_variaveis(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return list(DEFAULT_VARS)

    def _salvar_config_variaveis(self):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(self.variaveis_config, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("Erro ao salvar variáveis", str(e))

    def _criar_menu(self):
        menubar = tk.Menu(self.root)
        menu_arquivo = tk.Menu(menubar, tearoff=0)
        menu_arquivo.add_command(label="Novo Projeto Completo", command=self._novo_projeto)
        menu_arquivo.add_command(label="Abrir Projeto...", command=self._abrir_projeto)
        menu_arquivo.add_command(label="Salvar Projeto", command=self._salvar_projeto)
        menu_arquivo.add_separator()
        menu_arquivo.add_command(label="Sair", command=self.root.quit)
        menubar.add_cascade(label="Arquivo", menu=menu_arquivo)

        menu_config = tk.Menu(menubar, tearoff=0)
        menu_config.add_command(label="Configurar Variáveis...", command=self._janela_config_variaveis)
        menubar.add_cascade(label="Configurações", menu=menu_config)

        self.root.config(menu=menubar)

    def _criar_layout(self):
        frame_form = ttk.LabelFrame(self.root, text=" Cadastro do Dado de Mercado ", padding=10)
        frame_form.pack(fill="x", padx=10, pady=4)

        ttk.Label(frame_form, text="Informante:").grid(row=0, column=0, sticky="w")
        self.txt_informante = ttk.Entry(frame_form, width=22)
        self.txt_informante.grid(row=0, column=1, padx=4, pady=2)

        ttk.Label(frame_form, text="Telefone:").grid(row=0, column=2, sticky="w")
        self.txt_telefone = ttk.Entry(frame_form, width=22)
        self.txt_telefone.grid(row=0, column=3, padx=4, pady=2)

        ttk.Label(frame_form, text="Endereço/Logradouro:").grid(row=1, column=0, sticky="w")
        self.txt_endereco = ttk.Entry(frame_form, width=22)
        self.txt_endereco.grid(row=1, column=1, padx=4, pady=2)

        ttk.Label(frame_form, text="Bairro:").grid(row=1, column=2, sticky="w")
        self.txt_bairro = ttk.Entry(frame_form, width=22)
        self.txt_bairro.grid(row=1, column=3, padx=4, pady=2)

        ttk.Label(frame_form, text="Município:").grid(row=2, column=0, sticky="w")
        self.txt_municipio = ttk.Entry(frame_form, width=22)
        self.txt_municipio.grid(row=2, column=1, padx=4, pady=2)

        ttk.Label(frame_form, text="Valor Total (R$):").grid(row=2, column=2, sticky="w")
        self.txt_valor = ttk.Entry(frame_form, width=22)
        self.txt_valor.grid(row=2, column=3, padx=4, pady=2)

        ttk.Label(frame_form, text="Área Terreno:").grid(row=3, column=0, sticky="w")
        frame_area = ttk.Frame(frame_form)
        frame_area.grid(row=3, column=1, sticky="w", padx=4, pady=2)
        self.txt_area = ttk.Entry(frame_area, width=13)
        self.txt_area.pack(side="left")
        self.var_unidade = tk.StringVar(value="ha")
        self.cb_unidade = ttk.Combobox(frame_area, textvariable=self.var_unidade, values=["m²", "ha"], width=5, state="readonly")
        self.cb_unidade.pack(side="left", padx=2)

        ttk.Label(frame_form, text="Área Const. (m²):").grid(row=3, column=2, sticky="w")
        self.txt_area_const = ttk.Entry(frame_form, width=22)
        self.txt_area_const.grid(row=3, column=3, padx=4, pady=2)

        # Zona UTM e Coordenadas
        frame_coords = ttk.Frame(frame_form)
        frame_coords.grid(row=4, column=0, columnspan=4, sticky="w", pady=3)

        ttk.Label(frame_coords, text="Zona UTM:", foreground="#b22222", font=("Segoe UI", 9, "bold")).pack(side="left", padx=(0, 2))
        self.txt_zona = ttk.Entry(frame_coords, width=7)
        self.txt_zona.pack(side="left", padx=(0, 10))

        ttk.Label(frame_coords, text="Coord. E (m):").pack(side="left", padx=(0, 2))
        self.txt_coord_e = ttk.Entry(frame_coords, width=16)
        self.txt_coord_e.pack(side="left", padx=(0, 10))

        ttk.Label(frame_coords, text="Coord. S (m):").pack(side="left", padx=(0, 2))
        self.txt_coord_s = ttk.Entry(frame_coords, width=16)
        self.txt_coord_s.pack(side="left", padx=(0, 10))

        ttk.Label(frame_coords, text="Data:").pack(side="left", padx=(0, 2))
        self.txt_data = ttk.Entry(frame_coords, width=12)
        self.txt_data.insert(0, "03/09/2026")
        self.txt_data.pack(side="left")

        ttk.Label(frame_form, text="Link do Anúncio:").grid(row=5, column=0, sticky="w")
        self.txt_link = ttk.Entry(frame_form, width=65)
        self.txt_link.grid(row=5, column=1, columnspan=3, sticky="w", padx=4, pady=2)

        ttk.Label(frame_form, text="Foto 1 (Imóvel/Drone):").grid(row=6, column=0, sticky="w")
        self.txt_foto1 = ttk.Entry(frame_form, width=52)
        self.txt_foto1.grid(row=6, column=1, columnspan=2, sticky="w", padx=4, pady=2)
        ttk.Button(frame_form, text="Buscar", command=lambda: self._buscar_arquivo_foto(self.txt_foto1)).grid(row=6, column=3, sticky="w")

        ttk.Label(frame_form, text="Foto 2 (Print Anúncio):").grid(row=7, column=0, sticky="w")
        self.txt_foto2 = ttk.Entry(frame_form, width=52)
        self.txt_foto2.grid(row=7, column=1, columnspan=2, sticky="w", padx=4, pady=2)
        ttk.Button(frame_form, text="Buscar", command=lambda: self._buscar_arquivo_foto(self.txt_foto2)).grid(row=7, column=3, sticky="w")

        # Variáveis Dinâmicas
        self.frame_vars = ttk.LabelFrame(self.root, text=" Variáveis da Avaliação ", padding=10)
        self.frame_vars.pack(fill="x", padx=10, pady=4)
        self.widgets_dinamicos = {}

        frame_btn_cad = ttk.Frame(self.root, padding=4)
        frame_btn_cad.pack(fill="x", padx=10)

        self.btn_salvar_dado = ttk.Button(frame_btn_cad, text="➕ Adicionar Dado à Lista", command=self._adicionar_ou_salvar_dado)
        self.btn_salvar_dado.pack(side="left", padx=4)

        ttk.Button(frame_btn_cad, text="🧹 Novo Dado / Limpar Campos", command=self._limpar_formulario).pack(side="left", padx=4)

        self.btn_cancelar_edicao = ttk.Button(frame_btn_cad, text="✖ Cancelar Edição", command=self._limpar_formulario, state="disabled")
        self.btn_cancelar_edicao.pack(side="left", padx=4)

        ttk.Button(frame_btn_cad, text="⚙ Gerenciar Variáveis", command=self._janela_config_variaveis).pack(side="right", padx=4)

        # Tabela
        frame_tabela = ttk.LabelFrame(self.root, text=" Dados Cadastrados (Clique duplo para editar) ", padding=10)
        frame_tabela.pack(fill="both", expand=True, padx=10, pady=4)

        colunas = ("dado", "informante", "endereco", "municipio", "valor", "area", "unidade", "unitario")
        self.tree = ttk.Treeview(frame_tabela, columns=colunas, show="headings", height=7)
        self.tree.heading("dado", text="D.")
        self.tree.heading("informante", text="Informante")
        self.tree.heading("endereco", text="Endereço")
        self.tree.heading("municipio", text="Município")
        self.tree.heading("valor", text="Valor Total (R$)")
        self.tree.heading("area", text="Área")
        self.tree.heading("unidade", text="Unid.")
        self.tree.heading("unitario", text="Unitário (R$/un)")

        self.tree.column("dado", width=35, anchor="center")
        self.tree.column("unidade", width=55, anchor="center")
        self.tree.pack(fill="both", expand=True)

        self.tree.bind("<Double-1>", lambda event: self._carregar_para_edicao())

        frame_botoes_grid = ttk.Frame(frame_tabela)
        frame_botoes_grid.pack(fill="x", pady=4)
        ttk.Button(frame_botoes_grid, text="✏ Editar Selecionado", command=self._carregar_para_edicao).pack(side="left", padx=4)
        ttk.Button(frame_botoes_grid, text="🗑 Excluir Selecionado", command=self._excluir_dado).pack(side="left", padx=4)

        frame_acoes = ttk.Frame(self.root, padding=8)
        frame_acoes.pack(fill="x", padx=10, pady=4)

        ttk.Button(frame_acoes, text="📊 Exportar para Excel (SISDEA)", command=self._exportar_excel).pack(side="left", padx=8, expand=True, fill="x")
        ttk.Button(frame_acoes, text="📄 Exportar para Word (Fichas Técnicas)", command=self._exportar_word).pack(side="right", padx=8, expand=True, fill="x")

    def _atualizar_interface_variaveis(self):
        for w in self.frame_vars.winfo_children():
            w.destroy()
        self.widgets_dinamicos.clear()

        for idx, var in enumerate(self.variaveis_config):
            col = (idx % 2) * 2
            row = idx // 2

            lbl = ttk.Label(self.frame_vars, text=f"{var['nome']}:")
            lbl.grid(row=row, column=col, sticky="w", padx=4, pady=2)

            if var["tipo"] == "codigo" and var.get("opcoes"):
                cb = ttk.Combobox(self.frame_vars, values=var["opcoes"], width=24)
                if var["opcoes"]:
                    cb.set(var["opcoes"][0])
                cb.grid(row=row, column=col + 1, sticky="w", padx=4, pady=2)
                self.widgets_dinamicos[var["nome"]] = cb
            else:
                ent = ttk.Entry(self.frame_vars, width=26)
                ent.grid(row=row, column=col + 1, sticky="w", padx=4, pady=2)
                self.widgets_dinamicos[var["nome"]] = ent

    def _janela_config_variaveis(self):
        janela = tk.Toplevel(self.root)
        janela.title("Gerenciador de Variáveis da Pesquisa")
        janela.geometry("640x460")
        janela.transient(self.root)
        janela.grab_set()

        ttk.Label(janela, text="Defina as variáveis para o projeto (marque 'Ativar' para habilitar):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=6)

        frame_lista = ttk.Frame(janela, padding=6)
        frame_lista.pack(fill="both", expand=True)

        entradas_vars = []
        total_slots = max(6, len(self.variaveis_config) + 2)

        for i in range(total_slots):
            f_linha = ttk.Frame(frame_lista)
            f_linha.pack(fill="x", pady=2)

            ativa = i < len(self.variaveis_config)
            cfg = self.variaveis_config[i] if ativa else {"nome": "", "tipo": "texto", "opcoes": []}

            var_ativa = tk.BooleanVar(value=ativa)
            chk = ttk.Checkbutton(f_linha, text=f"Var {i+1}", variable=var_ativa)
            chk.grid(row=0, column=0, padx=2)

            e_nome = ttk.Entry(f_linha, width=16)
            e_nome.insert(0, cfg.get("nome", ""))
            e_nome.grid(row=0, column=1, padx=3)

            cb_tipo = ttk.Combobox(f_linha, values=["codigo", "numero", "texto"], width=8, state="readonly")
            cb_tipo.set(cfg.get("tipo", "texto"))
            cb_tipo.grid(row=0, column=2, padx=3)

            e_opcoes = ttk.Entry(f_linha, width=28)
            e_opcoes.insert(0, ", ".join(cfg.get("opcoes", [])))
            e_opcoes.grid(row=0, column=3, padx=3)

            entradas_vars.append((var_ativa, e_nome, cb_tipo, e_opcoes))

        def salvar_configs():
            novas_configs = []
            for var_ativa, e_nome, cb_tipo, e_opcoes in entradas_vars:
                if var_ativa.get() and e_nome.get().strip():
                    opts = [op.strip() for op in e_opcoes.get().split(",") if op.strip()]
                    novas_configs.append({
                        "nome": e_nome.get().strip(),
                        "tipo": cb_tipo.get(),
                        "opcoes": opts
                    })
            self.variaveis_config = novas_configs
            self._salvar_config_variaveis()
            self._atualizar_interface_variaveis()
            janela.destroy()
            messagebox.showinfo("Configurações", "Variáveis salvas com sucesso!")

        ttk.Button(janela, text="💾 Salvar Configurações", command=salvar_configs).pack(pady=8)

    def _buscar_arquivo_foto(self, entry_widget):
        caminho = filedialog.askopenfilename(filetypes=[("Imagens", "*.png;*.jpg;*.jpeg;*.webp")])
        if caminho:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, caminho)

    def _limpar_formulario(self):
        self.item_em_edicao = None
        self.txt_informante.delete(0, tk.END)
        self.txt_telefone.delete(0, tk.END)
        self.txt_endereco.delete(0, tk.END)
        self.txt_bairro.delete(0, tk.END)
        self.txt_municipio.delete(0, tk.END)
        self.txt_valor.delete(0, tk.END)
        self.txt_area.delete(0, tk.END)
        self.txt_area_const.delete(0, tk.END)
        self.txt_zona.delete(0, tk.END)
        self.txt_coord_e.delete(0, tk.END)
        self.txt_coord_s.delete(0, tk.END)
        self.txt_data.delete(0, tk.END)
        self.txt_data.insert(0, "03/09/2026")
        self.txt_link.delete(0, tk.END)
        self.txt_foto1.delete(0, tk.END)
        self.txt_foto2.delete(0, tk.END)

        for _, widget in self.widgets_dinamicos.items():
            if isinstance(widget, ttk.Combobox):
                vals = widget.cget("values")
                if vals:
                    widget.set(vals[0])
            else:
                widget.delete(0, tk.END)

        self.btn_salvar_dado.config(text="➕ Adicionar Dado à Lista")
        self.btn_cancelar_edicao.config(state="disabled")

    def _adicionar_ou_salvar_dado(self):
        try:
            valor_total = converter_para_float(self.txt_valor.get())
            area_num = converter_para_float(self.txt_area.get())
            area_const = converter_para_float(self.txt_area_const.get())
            unidade = self.var_unidade.get()
            unitario = valor_total / area_num if area_num > 0 else 0.0

            dado_num = self.item_em_edicao["D."] if self.item_em_edicao else (len(self.dados_pesquisas) + 1)

            # Limpa qualquer "m E" ou "m S" pré-existente
            coord_e_pura = limpar_sufixo_coord(self.txt_coord_e.get())
            coord_s_pura = limpar_sufixo_coord(self.txt_coord_s.get())

            registro = {
                "D.": dado_num,
                "Informante": self.txt_informante.get().strip(),
                "Telefone": self.txt_telefone.get().strip(),
                "Endereço": self.txt_endereco.get().strip(),
                "Bairro": self.txt_bairro.get().strip(),
                "Município": self.txt_municipio.get().strip(),
                "Valor Total (R$)": valor_total,
                f"Área Terreno ({unidade})": area_num,
                "Área Construída (m²)": area_const,
                f"Unitário (R$/{unidade})": unitario,
                "Localização": "Rural" if unidade == "ha" else "Urbana",
                "Zona UTM": self.txt_zona.get().strip(),
                "Coord. E (m)": coord_e_pura,
                "Coord. S (m)": coord_s_pura,
                "Data": self.txt_data.get().strip(),
                "Link": self.txt_link.get().strip(),
                "Foto1": self.txt_foto1.get().strip(),
                "Foto2": self.txt_foto2.get().strip(),
                "Unidade": unidade,
                "VariaveisExtras": {}
            }

            for nome, widget in self.widgets_dinamicos.items():
                val = widget.get().strip()
                registro[nome] = val
                registro["VariaveisExtras"][nome] = val

            if self.item_em_edicao:
                idx = next(i for i, d in enumerate(self.dados_pesquisas) if d["D."] == dado_num)
                self.dados_pesquisas[idx] = registro
                messagebox.showinfo("Atualização", f"Pesquisa {dado_num:02d} atualizada com sucesso!")
            else:
                self.dados_pesquisas.append(registro)

            self._recarregar_grid()
            self._limpar_formulario()
        except Exception as e:
            messagebox.showerror("Erro de Preenchimento", f"Verifique os campos numéricos: {e}")

    def _recarregar_grid(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        for dado in self.dados_pesquisas:
            un = dado.get("Unidade", "ha")
            casas = 4 if un == "ha" else 2
            v_total = dado.get("Valor Total (R$)", 0.0)
            a_total = dado.get(f"Área Terreno ({un})", 0.0)
            u_unit = dado.get(f"Unitário (R$/{un})", 0.0)

            self.tree.insert("", "end", values=(
                dado["D."], dado["Informante"], dado["Endereço"],
                dado["Município"], f"R$ {formatar_moeda_br(v_total)}",
                f"{formatar_numero_br(a_total, casas)}",
                un, f"R$ {formatar_moeda_br(u_unit)}"
            ))

    def _carregar_para_edicao(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("Aviso", "Selecione um dado na tabela para editar.")
            return

        item_val = self.tree.item(sel[0])["values"]
        dado_id = item_val[0]
        dado = next((d for d in self.dados_pesquisas if d["D."] == dado_id), None)
        if not dado:
            return

        self.item_em_edicao = dado
        self.btn_salvar_dado.config(text=f"💾 Atualizar Dado {dado_id}")
        self.btn_cancelar_edicao.config(state="normal")

        self.txt_informante.delete(0, tk.END)
        self.txt_informante.insert(0, dado.get("Informante", ""))

        self.txt_telefone.delete(0, tk.END)
        self.txt_telefone.insert(0, dado.get("Telefone", ""))

        self.txt_endereco.delete(0, tk.END)
        self.txt_endereco.insert(0, dado.get("Endereço", ""))

        self.txt_bairro.delete(0, tk.END)
        self.txt_bairro.insert(0, dado.get("Bairro", ""))

        self.txt_municipio.delete(0, tk.END)
        self.txt_municipio.insert(0, dado.get("Município", ""))

        # Carrega o valor total formatado em padrão brasileiro (evita ponto decimal do python)
        v_total = dado.get("Valor Total (R$)", 0.0)
        self.txt_valor.delete(0, tk.END)
        self.txt_valor.insert(0, formatar_moeda_br(v_total))

        un = dado.get("Unidade", "ha")
        self.var_unidade.set(un)

        # Carrega área com formato brasileiro
        a_total = dado.get(f"Área Terreno ({un})", 0.0)
        casas = 4 if un == "ha" else 2
        self.txt_area.delete(0, tk.END)
        self.txt_area.insert(0, formatar_numero_br(a_total, casas))

        a_const = dado.get("Área Construída (m²)", 0.0)
        self.txt_area_const.delete(0, tk.END)
        if a_const > 0:
            self.txt_area_const.insert(0, formatar_numero_br(a_const, 2))

        self.txt_zona.delete(0, tk.END)
        self.txt_zona.insert(0, dado.get("Zona UTM", ""))

        self.txt_coord_e.delete(0, tk.END)
        self.txt_coord_e.insert(0, dado.get("Coord. E (m)", ""))

        self.txt_coord_s.delete(0, tk.END)
        self.txt_coord_s.insert(0, dado.get("Coord. S (m)", ""))

        self.txt_data.delete(0, tk.END)
        self.txt_data.insert(0, dado.get("Data", ""))

        self.txt_link.delete(0, tk.END)
        self.txt_link.insert(0, dado.get("Link", ""))

        self.txt_foto1.delete(0, tk.END)
        self.txt_foto1.insert(0, dado.get("Foto1", ""))

        self.txt_foto2.delete(0, tk.END)
        self.txt_foto2.insert(0, dado.get("Foto2", ""))

        extras = dado.get("VariaveisExtras", {})
        for nome, widget in self.widgets_dinamicos.items():
            val = extras.get(nome, dado.get(nome, ""))
            if isinstance(widget, ttk.Combobox):
                widget.set(val)
            else:
                widget.delete(0, tk.END)
                widget.insert(0, str(val))

    def _excluir_dado(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("Aviso", "Selecione um dado na tabela para excluir.")
            return

        item_val = self.tree.item(sel[0])["values"]
        dado_id = item_val[0]
        if messagebox.askyesno("Confirmar Exclusão", f"Deseja realmente excluir o Dado {dado_id}?"):
            self.dados_pesquisas = [d for d in self.dados_pesquisas if d["D."] != dado_id]
            for i, d in enumerate(self.dados_pesquisas):
                d["D."] = i + 1
            self._recarregar_grid()
            self._limpar_formulario()

    def _novo_projeto(self):
        if messagebox.askyesno("Novo Projeto", "Deseja iniciar um novo projeto limpo? Todos os dados não salvos serão descartados."):
            self.dados_pesquisas = []
            self._recarregar_grid()
            self._limpar_formulario()

    def _salvar_projeto(self):
        caminho = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("Projeto Pesquisa (*.json)", "*.json")])
        if not caminho:
            return
        dados_salvar = {
            "config_variaveis": self.variaveis_config,
            "pesquisas": self.dados_pesquisas
        }
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(dados_salvar, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("Sucesso", "Projeto salvo com sucesso!")

    def _abrir_projeto(self):
        caminho = filedialog.askopenfilename(filetypes=[("Projeto Pesquisa (*.json)", "*.json")])
        if not caminho:
            return
        try:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = json.load(f)
            self.variaveis_config = conteudo.get("config_variaveis", self.variaveis_config)
            self.dados_pesquisas = conteudo.get("pesquisas", [])
            self._atualizar_interface_variaveis()
            self._recarregar_grid()
            self._limpar_formulario()
            messagebox.showinfo("Sucesso", "Projeto carregado com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro ao abrir", f"Não foi possível abrir o arquivo: {e}")

    def _baixar_imagem(self, caminho_ou_url):
        if not caminho_ou_url:
            return None
        try:
            if caminho_ou_url.startswith("http://") or caminho_ou_url.startswith("https://"):
                url = caminho_ou_url
                if "drive.google.com" in url and "id=" in url:
                    file_id = re.search(r"id=([a-zA-Z0-9_-]+)", url).group(1)
                    url = f"https://drive.google.com/uc?export=download&id={file_id}"
                resp = requests.get(url, timeout=10)
                if resp.status_code == 200:
                    return io.BytesIO(resp.content)
            elif os.path.exists(caminho_ou_url):
                return caminho_ou_url
        except Exception:
            return None
        return None

    def _exportar_excel(self):
        if not self.dados_pesquisas:
            messagebox.showwarning("Aviso", "Nenhum dado cadastrado.")
            return

        caminho = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")])
        if not caminho:
            return

        df = pd.DataFrame(self.dados_pesquisas)
        df_sisdea = df.drop(columns=["Foto1", "Foto2", "Unidade", "VariaveisExtras"], errors="ignore")

        with pd.ExcelWriter(caminho, engine="openpyxl") as writer:
            df_sisdea.to_excel(writer, sheet_name="Ficha de Pesquisa", index=False)

        messagebox.showinfo("Concluído", "Planilha Excel para SISDEA exportada com sucesso!")

    def _definir_bordas_tabela(self, table):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def _exportar_word(self):
        if not self.dados_pesquisas:
            messagebox.showwarning("Aviso", "Nenhum dado cadastrado.")
            return

        caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if not caminho:
            return

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.35)
            section.bottom_margin = Inches(0.35)
            section.left_margin = Inches(0.45)
            section.right_margin = Inches(0.45)

        total_dados = len(self.dados_pesquisas)
        for i in range(0, total_dados, 2):
            if i > 0:
                doc.add_page_break()

            p_tit = doc.add_paragraph()
            p_tit.paragraph_format.space_before = Pt(0)
            p_tit.paragraph_format.space_after = Pt(2)
            r_tit = p_tit.add_run("PESQUISA DE MERCADO - MEMÓRIA DE CÁLCULO")
            r_tit.bold = True
            r_tit.font.name = "Arial"
            r_tit.font.size = Pt(10)

            p_bar = doc.add_paragraph()
            p_bar.paragraph_format.space_before = Pt(0)
            p_bar.paragraph_format.space_after = Pt(4)
            pBrd = parse_xml(
                f'<w:pBrd {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="245D8C"/>'
                f'</w:pBrd>'
            )
            p_bar._p.get_or_add_pPr().append(pBrd)

            tabela = doc.add_table(rows=0, cols=2)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabela.autofit = False
            self._definir_bordas_tabela(tabela)

            lote = self.dados_pesquisas[i:i+2]
            for dado in lote:
                un = dado.get("Unidade", "ha")
                row = tabela.add_row()
                celula_dados, celula_fotos = row.cells[0], row.cells[1]
                celula_dados.width = Inches(3.7)
                celula_fotos.width = Inches(3.7)
                celula_dados.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                celula_fotos.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Título da Pesquisa centralizado no topo
                p_num = celula_dados.paragraphs[0]
                p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_num.paragraph_format.space_before = Pt(0)
                p_num.paragraph_format.space_after = Pt(4)
                r_num = p_num.add_run(f"Pesquisa – {dado['D.']:02d}")
                r_num.bold = True
                r_num.font.name = "Arial"
                r_num.font.size = Pt(10)

                p_dados = celula_dados.add_paragraph()
                p_dados.paragraph_format.line_spacing = 1.10
                p_dados.paragraph_format.space_before = Pt(0)
                p_dados.paragraph_format.space_after = Pt(0)

                def add_f_line(p, label, val):
                    r1 = p.add_run(label)
                    r1.bold = True
                    r1.font.name = "Arial"
                    r1.font.size = Pt(10)
                    r2 = p.add_run(f" {val}\n")
                    r2.font.name = "Arial"
                    r2.font.size = Pt(10)

                add_f_line(p_dados, "Logradouro:", dado.get("Endereço", ""))
                if dado.get("Bairro"):
                    add_f_line(p_dados, "Bairro:", dado.get("Bairro", ""))
                add_f_line(p_dados, "Município:", dado.get("Município", ""))
                add_f_line(p_dados, "Contato:", f"{dado.get('Telefone', '')} - {dado.get('Informante', '')}")
                add_f_line(p_dados, "Link:", dado.get("Link", ""))
                p_dados.add_run("\n")

                area_val = dado.get(f"Área Terreno ({un})", 0.0)
                casas_area = 4 if un == "ha" else 2
                add_f_line(p_dados, "Área Terreno:", f"{formatar_numero_br(area_val, casas_area)} {un}")

                const_val = dado.get("Área Construída (m²)", 0.0)
                add_f_line(p_dados, "Área Construída:", f"{formatar_numero_br(const_val, 2)} m²")

                v_total = dado.get("Valor Total (R$)", 0.0)
                add_f_line(p_dados, "Valor da Oferta:", f"R$ {formatar_moeda_br(v_total)}")

                u_val = dado.get(f"Unitário (R$/{un})", 0.0)
                add_f_line(p_dados, f"Valor Unitário/{un}:", f"R$ {formatar_moeda_br(u_val)}/{un}")
                p_dados.add_run("\n")

                for v_cfg in self.variaveis_config:
                    v_nome = v_cfg["nome"]
                    v_val = dado.get("VariaveisExtras", {}).get(v_nome, dado.get(v_nome, ""))
                    add_f_line(p_dados, f"{v_nome}:", v_val)

                # Formatação das Coordenadas Geográficas
                zona_str = f"{dado.get('Zona UTM', '').strip()} " if dado.get('Zona UTM') else ""
                coord_e = limpar_sufixo_coord(dado.get('Coord. E (m)', ''))
                coord_s = limpar_sufixo_coord(dado.get('Coord. S (m)', ''))
                coord_texto = f"{zona_str}{coord_e} m E / {coord_s} m S" if coord_e or coord_s else ""

                add_f_line(p_dados, "Coordenadas Geográfica:", coord_texto)
                add_f_line(p_dados, "Localização:", dado.get("Localização", "Rural"))
                p_dados.add_run("\n")
                add_f_line(p_dados, "Data:", dado.get("Data", ""))

                # Inserção das Fotos
                img1 = self._baixar_imagem(dado.get("Foto1"))
                img2 = self._baixar_imagem(dado.get("Foto2"))

                p_foto = celula_fotos.paragraphs[0]
                p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_foto.paragraph_format.space_before = Pt(2)
                p_foto.paragraph_format.space_after = Pt(2)

                if img1 and img2:
                    try:
                        p_foto.add_run().add_picture(img1, width=Inches(2.75))
                    except Exception:
                        p_foto.add_run("[ Erro na Foto 1 ]\n")

                    p_foto2 = celula_fotos.add_paragraph()
                    p_foto2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_foto2.paragraph_format.space_before = Pt(2)
                    p_foto2.paragraph_format.space_after = Pt(2)
                    try:
                        p_foto2.add_run().add_picture(img2, width=Inches(2.75))
                    except Exception:
                        p_foto2.add_run("[ Erro na Foto 2 ]")
                elif img1:
                    try:
                        p_foto.add_run().add_picture(img1, width=Inches(2.75))
                    except Exception:
                        p_foto.add_run("[ Erro na Foto 1 ]")
                elif img2:
                    try:
                        p_foto.add_run().add_picture(img2, width=Inches(2.75))
                    except Exception:
                        p_foto.add_run("[ Erro na Foto 2 ]")
                else:
                    r_vazio = p_foto.add_run("[ Sem fotos anexadas ]")
                    r_vazio.font.name = "Arial"
                    r_vazio.font.size = Pt(10)

        doc.save(caminho)
        messagebox.showinfo("Exportação Concluída", "Fichas de Pesquisa no Word geradas com sucesso!")

if __name__ == "__main__":
    root = tk.Tk()
    app = AppPesquisaMercado(root)
    root.mainloop()
